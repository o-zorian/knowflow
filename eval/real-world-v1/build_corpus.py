from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
CORPUS = ROOT / "corpus"
DATASET = ROOT / "real-world-v1.jsonl"
MANIFEST = ROOT / "documents.json"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            set_cell_margins(cell)


def set_run_font(run, size=11, bold=False, color="222222"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, 9, color="6B7280")


def configure_docx(document, running_label):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Title", 25, 0, 8, "0B2545"),
        ("Subtitle", 12, 0, 18, "4B5563"),
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Title":
            p_pr = style._element.get_or_add_pPr()
            borders = p_pr.find(qn("w:pBdr"))
            if borders is not None:
                p_pr.remove(borders)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(running_label), 9, bold=True, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("Real-world-v1 | Page "), 9, color="6B7280")
    add_page_field(footer)


def add_docx_title(document, title, subtitle, metadata):
    p = document.add_paragraph(style="Title")
    set_run_font(p.add_run(title), 25, bold=True, color="0B2545")
    p = document.add_paragraph(style="Subtitle")
    set_run_font(p.add_run(subtitle), 12, color="4B5563")
    for label, value in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(label + "："), 10.5, bold=True, color="374151")
        set_run_font(p.add_run(value), 10.5, color="374151")
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(document, text, level=1):
    p = document.add_paragraph(style=f"Heading {level}")
    set_run_font(p.add_run(text), {1: 16, 2: 13, 3: 12}[level], bold=True, color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level])


def add_body(document, text, bold_lead=None):
    p = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), 11, bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]), 11)
    else:
        set_run_font(p.add_run(text), 11)
    return p


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        cell._tc.tcPr[-1].set(qn("w:fill"), "E8EEF5")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.keep_with_next = True
        set_run_font(cell.paragraphs[0].add_run(header), 9.5, bold=True, color="0B2545")
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(cells[index].paragraphs[0].add_run(str(value)), 9.5)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def build_helios_current(path):
    doc = Document()
    configure_docx(doc, "Helios Mobility | Travel Policy v3.2")
    add_docx_title(doc, "Helios 员工差旅政策", "版本 3.2（当前有效版）", [("生效日期", "2026-06-01"), ("替代版本", "v2.8 及此前版本"), ("政策负责人", "财务运营与人员体验部")])
    add_heading(doc, "1. 适用范围与版本优先级")
    add_body(doc, "本政策适用于 Helios 中国区员工因公司业务产生的境内及国际差旅。版本 3.2 自 2026 年 6 月 1 日生效，并明确替代版本 2.8；归档政策只能用于历史报销审计，不能用于新的预订或审批。员工看到不同版本的金额或舱等规则时，必须以本文件标注的当前有效版为准。")
    add_heading(doc, "2. 交通舱等与提前预订")
    add_body(doc, "单程计划飞行时间不足 6 小时应预订经济舱；6 小时至 10 小时可预订高级经济舱；超过 10 小时仅在业务副总裁书面批准后可预订商务舱。所有机票原则上应在出发前至少 14 个自然日完成预订，临时客户事件可以缩短窗口，但必须在报销单中填写原因。")
    add_body(doc, "铁路出行以二等座为默认标准。项目负责人不能把“行程紧急”直接等同于商务舱授权；时间窗口例外与舱等例外是两个独立审批事项，分别留下证据。")
    add_heading(doc, "3. 住宿与会议例外")
    add_table(doc, ["城市", "普通每晚上限", "指定会议例外", "币种"], [("上海", "900", "最高上浮 20%", "人民币"), ("北京", "850", "最高上浮 20%", "人民币"), ("深圳", "800", "最高上浮 20%", "人民币"), ("其他境内城市", "650", "最高上浮 15%", "人民币")], [2160, 2160, 2880, 2160])
    add_body(doc, "上海普通住宿上限为每晚 900 元。参加公司指定会议且会场酒店为唯一合理选择时，上限最多上浮 20%，因此上海的会议例外上限为每晚 1,080 元；例外不自动覆盖迷你吧、升级房型或陪同人员费用。")
    add_heading(doc, "4. 餐费、杂费与审批")
    add_body(doc, "境内差旅餐费按每日 300 元封顶，必要杂费按每日 80 元封顶。单次行程预计总额不超过 8,000 元由直属经理批准；8,001 至 25,000 元需要部门负责人批准；超过 25,000 元还需要财务业务伙伴复核。餐费与杂费是不同额度，不能相互挪用。")
    add_heading(doc, "5. 干扰信息与审计说明")
    add_body(doc, "2025 年试行的“绿色出行积分”曾允许用积分抵扣酒店超额，但该试行已在 2026 年 3 月结束，不属于版本 3.2 的报销依据。差旅帮助台可能在邮件中引用旧示例，审批人仍需回到当前政策核对日期、金额与权限。")
    doc.save(path)


def build_helios_archived(path):
    doc = Document()
    configure_docx(doc, "Helios Mobility | ARCHIVED Travel Policy v2.8")
    add_docx_title(doc, "Helios 员工差旅政策（归档）", "版本 2.8 - 已被 v3.2 替代", [("历史生效期", "2025-01-15 至 2026-05-31"), ("状态", "ARCHIVED / 不得用于新预订"), ("用途", "仅供历史报销审计")])
    add_heading(doc, "1. 归档警示")
    add_body(doc, "本文件自 2026 年 6 月 1 日起不再有效。任何在该日期之后发生的新预订必须查阅版本 3.2；本文件中的较低住宿上限、较短预订窗口和较宽松商务舱门槛都是历史规则。")
    add_heading(doc, "2. 历史舱等规则")
    add_body(doc, "版本 2.8 规定：单程不足 5 小时使用经济舱，5 小时至 8 小时可使用高级经济舱，超过 8 小时在部门负责人批准后可使用商务舱。机票应至少提前 7 个自然日预订。上述规则仅用于判断 2026 年 5 月 31 日及以前已完成行程的合规性。")
    add_heading(doc, "3. 历史住宿与日补贴")
    add_table(doc, ["城市", "历史住宿上限", "历史餐费", "币种"], [("上海", "760", "260/日", "人民币"), ("北京", "720", "260/日", "人民币"), ("深圳", "700", "260/日", "人民币"), ("其他境内城市", "580", "260/日", "人民币")], [2340, 2340, 2340, 2340])
    add_body(doc, "归档版本中，上海住宿上限为每晚 760 元，境内餐费为每日 260 元。与当前版本相比，上海普通住宿上限后来增加了 140 元，餐费增加了 40 元。")
    add_heading(doc, "4. 审计边界")
    add_body(doc, "审计人员必须按费用发生日期选择政策版本，不得因为归档文件搜索排名更高就用于当前行程。版本号相似并不代表规则可互换，尤其要区分“超过 8 小时可申请商务舱”和当前版本“超过 10 小时且需业务副总裁批准”。")
    doc.save(path)


def build_atlas(path):
    doc = Document()
    configure_docx(doc, "Atlas Support | Incident Response Playbook 2026.4")
    add_docx_title(doc, "Atlas 客户支持事件响应手册", "版本 2026.4 - 严重度、优先级与升级路径", [("生效日期", "2026-04-20"), ("所有者", "客户可靠性办公室"), ("复审周期", "每 180 天")])
    add_heading(doc, "1. 严重度与优先级不是同一概念")
    add_body(doc, "严重度 S1/S2/S3 描述客户影响和系统状态；优先级 P1/P2/P3 描述内部工作排序。一个影响单个战略客户但有替代方案的问题可以是 S2、P1；因此客服不得把 P1 自动说成 S1，也不得用优先级替代对可用性和数据风险的判断。")
    add_heading(doc, "2. 响应目标")
    add_table(doc, ["严重度", "首次响应", "状态更新", "事件指挥官", "高管通知"], [("S1", "15 分钟", "每 30 分钟", "立即指定", "20 分钟内"), ("S2", "1 小时", "每 2 小时", "按需指定", "通常不要求"), ("S3", "1 个工作日", "重大变化时", "不要求", "不要求")], [1400, 1700, 1700, 2100, 2460])
    add_body(doc, "S1 事件的首次响应目标是 15 分钟，随后每 30 分钟更新一次；事件指挥官应立即指定，高管通知必须在确认 S1 后 20 分钟内完成。S2 的首次响应目标是 1 小时，更新间隔为 2 小时。")
    add_heading(doc, "3. 跨团队升级与交接")
    add_body(doc, "值班支持工程师在确认 S1 后先建立事件频道并指定临时记录员，再由事件指挥官在 10 分钟内确认客户沟通负责人。工程团队修复服务并不自动结束事件：客户沟通负责人必须确认外部状态页已更新，记录员必须链接时间线，最后由事件指挥官宣布恢复。")
    add_heading(doc, "4. 数据风险与相似术语")
    add_body(doc, "“数据延迟”表示数据最终会到达但超过承诺时间；“数据丢失”表示记录不可恢复或需要从备份重建。前者通常从 S2 评估，后者只要涉及生产客户数据就从 S1 评估。术语相近，但升级门槛完全不同。")
    add_heading(doc, "5. 退款与合同边界")
    add_body(doc, "本手册不规定客户退款金额、服务信用额度或合同赔偿上限。支持团队只能记录影响时段并转交客户成功与法务，不能从本手册推断具体退款比例。")
    doc.save(path)


def pdf_styles():
    # A standalone OpenType font produces a reliable ToUnicode map for PDF
    # extractors.  The TTC YaHei files render correctly but some Linux PDF
    # parsers decode line-ending glyphs as unrelated CJK characters, which is
    # unacceptable for a retrieval-quality corpus.
    # Use a static TTF; ReportLab's subsetter cannot emit a correct ToUnicode
    # map from the variable NotoSansSC font on Windows.
    chinese_font = r"C:\Windows\Fonts\simhei.ttf"
    pdfmetrics.registerFont(TTFont("NotoSansSC", chinese_font))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("CNTitle", parent=styles["Title"], fontName="NotoSansSC", fontSize=23, leading=30, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER, spaceAfter=14),
        "subtitle": ParagraphStyle("CNSubtitle", parent=styles["Normal"], fontName="NotoSansSC", fontSize=11, leading=16, textColor=colors.HexColor("#4B5563"), alignment=TA_CENTER, spaceAfter=18),
        "h1": ParagraphStyle("CNH1", parent=styles["Heading1"], fontName="NotoSansSC", fontSize=15, leading=21, textColor=colors.HexColor("#2E74B5"), spaceBefore=14, spaceAfter=8),
        "h2": ParagraphStyle("CNH2", parent=styles["Heading2"], fontName="NotoSansSC", fontSize=12, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=10, spaceAfter=6),
        "body": ParagraphStyle("CNBody", parent=styles["BodyText"], fontName="NotoSansSC", fontSize=10.5, leading=17, textColor=colors.HexColor("#222222"), alignment=TA_LEFT, spaceAfter=8),
        "small": ParagraphStyle("CNSmall", parent=styles["BodyText"], fontName="NotoSansSC", fontSize=8.5, leading=12, textColor=colors.HexColor("#555555"), spaceAfter=4),
    }


def pdf_footer(canvas, document, label):
    canvas.saveState()
    canvas.setFont("NotoSansSC", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(inch, 0.48 * inch, label)
    canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Real-world-v1 | Page {document.page}")
    canvas.restoreState()


def build_pdf(path, title, subtitle, label, sections):
    styles = pdf_styles()
    story = [Spacer(1, 0.25 * inch), Paragraph(title, styles["title"]), Paragraph(subtitle, styles["subtitle"])]
    for section in sections:
        kind = section[0]
        if kind == "pagebreak":
            story.append(PageBreak())
        elif kind == "h1":
            story.append(Paragraph(section[1], styles["h1"]))
        elif kind == "h2":
            story.append(Paragraph(section[1], styles["h2"]))
        elif kind == "p":
            story.append(Paragraph(section[1], styles["body"]))
        elif kind == "table":
            data, widths = section[1], section[2]
            wrapped = [[Paragraph(str(value), styles["small"]) for value in row] for row in data]
            table = Table(wrapped, colWidths=widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("FONTNAME", (0, 0), (-1, 0), "NotoSansSC"),
                ("FONTNAME", (0, 1), (-1, -1), "NotoSansSC"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAB4C3")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.extend([table, Spacer(1, 8)])
    document = SimpleDocTemplate(str(path), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=0.85 * inch, bottomMargin=0.75 * inch, title=title, author="KnowFlow Real-world Evaluation")
    callback = lambda canvas, doc: pdf_footer(canvas, doc, label)
    document.build(story, onFirstPage=callback, onLaterPages=callback)


def build_pdfs():
    build_pdf(CORPUS / "aurora-transit-operations-v2.1.pdf", "Aurora 城市交通运营手册", "版本 2.1 | 当前有效 | 2026-04-01", "Aurora Operations v2.1", [
        ("h1", "1. 版本与适用边界"),
        ("p", "版本 2.1 自 2026 年 4 月 1 日生效，替代版本 1.9。手册覆盖 Aurora 轨道交通日常运营、服务中断、恶劣天气和失物管理。2024 年“蓝卡快速放行”试点已结束；它出现在历史附录中只是为了审计，不能作为当前运营指令。"),
        ("h1", "2. 事件等级与响应时限"),
        ("table", [["等级", "名称", "控制中心确认", "主要动作"], ["P0", "Safety Stop 安全停车", "立即", "停止相关区段运行"], ["P1", "Service Critical 服务关键", "10 分钟", "25 分钟内安排接驳车"], ["P2", "Degraded 降级运行", "30 分钟", "4 小时内恢复或给出计划"]], [0.65*inch, 1.55*inch, 1.35*inch, 2.95*inch]),
        ("p", "P1 服务关键事件要求控制中心在 10 分钟内确认，并在 25 分钟内安排替代接驳车。P2 降级运行的确认目标是 30 分钟，恢复目标为 4 小时。"),
        ("h2", "2.1 相似术语"),
        ("p", "“Safety Stop 安全停车”是基于安全风险的强制停止，对应 P0；“Service Pause 服务暂停”是为恢复班次或调度间隔而进行的短时计划动作，本身不等于 P0。两者都可能让列车静止，但触发条件、通报对象和恢复授权不同。"),
        ("h1", "3. 班次与备用运力"),
        ("table", [["线路", "高峰间隔", "非高峰间隔", "最低备用车辆比例"], ["A 线", "6 分钟", "10 分钟", "12%"], ["B 线", "8 分钟", "12 分钟", "10%"], ["机场快线", "15 分钟", "20 分钟", "15%"]], [1.25*inch, 1.5*inch, 1.65*inch, 2.1*inch]),
        ("p", "机场快线的高峰计划间隔为 15 分钟，非高峰为 20 分钟，最低备用车辆比例为 15%。A 线的最低备用比例是 12%，不能与机场快线混淆。"),
        ("pagebreak",),
        ("h1", "4. 恶劣天气跨章节处置"),
        ("p", "当气象台发布暴雨红色预警，且高架区段连续 10 分钟测得平均风速超过 24 米/秒时，值班主任应暂停受影响高架区段，并按 P1 流程通知控制中心。暂停后还必须在 15 分钟内发布乘客公告，并协调地面接驳；只满足红色预警而未达到风速持续条件时，应加强监测但不自动停运。"),
        ("h1", "5. 失物分类与保管"),
        ("p", "普通失物使用登记代码 AUR-17，保管 90 天后按批准流程处置。含裸露或损坏锂电池的物品属于危险品，必须在 24 小时内转交安全团队，不能按普通失物保管 90 天。身份证件则在登记后 48 小时内移交公安联络点。"),
        ("h1", "6. 干扰内容：已终止试点"),
        ("p", "蓝卡试点曾把部分 P2 事件的确认时间缩短到 20 分钟，但试点仅在 2024 年 7 月至 9 月运行，版本 2.1 明确恢复正式的 30 分钟目标。检索结果出现“20 分钟”时必须核对它是否来自已终止试点。"),
    ])
    build_pdf(CORPUS / "aurora-maintenance-bulletin-2026-rB.pdf", "Aurora NX-7 维护公告", "公告 2026-17 | Revision B | 2026-05-18", "Aurora Bulletin 2026-17 Rev B", [
        ("h1", "1. 修订状态"),
        ("p", "Revision B 是公告 2026-17 的当前修订版，发布日期为 2026 年 5 月 18 日。它把 NX-7 车门执行器安装扭矩从 Revision A 的 38 N·m 更正为 42 N·m；所有未关闭的工单必须采用 42 N·m，并在记录中注明 Revision B。"),
        ("h1", "2. 检查周期与零件号"),
        ("p", "NX-7 车门执行器应每 18,000 公里或每 120 天检查一次，以先到者为准。导向滚轮零件号是 AX-440，位置传感器零件号是 AX-404；两个编号数字相似，但不得互换。"),
        ("table", [["项目", "正确零件号", "检查要点", "处置"], ["导向滚轮", "AX-440", "磨耗与偏心", "超过 0.8 mm 更换"], ["位置传感器", "AX-404", "信号漂移", "超过 2.0% 校准"], ["执行器支架", "AX-414", "裂纹", "任何裂纹立即更换"]], [1.25*inch, 1.25*inch, 2.0*inch, 2.0*inch]),
        ("h1", "3. 车门隔离与运营等级"),
        ("p", "单个车门被隔离且相邻车门可用时，车辆可按 P2 降级运行；同一车厢两个相邻车门不可用时，应升级为 P1，并执行运营手册中的 P1 确认与接驳要求。该规则需要把本公告的技术状态与运营手册的响应时限结合起来。"),
        ("pagebreak",),
        ("h1", "4. 维护窗口"),
        ("table", [["地点", "日期", "停电窗口", "受影响车辆"], ["Central Depot", "2026-06-08", "01:10-03:40", "NX-7 021-036"], ["East Yard", "2026-06-10", "00:30-02:00", "NX-7 041-048"], ["Airport Siding", "2026-06-12", "02:15-04:15", "NX-7 051-054"]], [1.45*inch, 1.35*inch, 1.65*inch, 2.05*inch]),
        ("p", "Central Depot 的维护停电窗口为 2026 年 6 月 8 日 01:10 至 03:40，覆盖 NX-7 021 至 036。East Yard 的窗口不是同一天，也不覆盖该车辆范围。"),
        ("h1", "5. 完工证据"),
        ("p", "关闭工单前必须上传扭矩扳手读数照片、执行器序列号和双人复核签名。只有文字备注“已检查”不能作为完工证据；缺任一项时工单保持 blocked，而不是 completed。"),
    ])
    build_pdf(CORPUS / "northwind-energy-q2-2026.pdf", "Northwind 园区能源绩效报告", "2026 Q2 | Final 版 | 财务与可持续发展联合发布", "Northwind Energy Q2 2026 Final", [
        ("h1", "1. 报告口径与版本"),
        ("p", "本报告为 2026 年第二季度 Final 版，计量窗口是 4 月 1 日至 6 月 30 日。4 月发布的 provisional 工作表曾把 Harbor 站点记录为 421 MWh，但最终对账确认其实际用电为 418 MWh；所有季度比较应使用 418 MWh。"),
        ("h1", "2. 站点能源表"),
        ("table", [["站点", "实际用电 MWh", "目标 MWh", "差异 MWh", "排放 tCO2e"], ["Harbor", "418", "430", "-12", "96.2"], ["Ridge", "367", "350", "+17", "81.4"], ["Delta", "512", "500", "+12", "119.8"], ["合计", "1,297", "1,280", "+17", "297.4"]], [1.2*inch, 1.35*inch, 1.2*inch, 1.2*inch, 1.55*inch]),
        ("p", "三个站点实际用电合计 1,297 MWh，目标合计 1,280 MWh，超出 17 MWh，即 1.33%。Ridge 实际用电 367 MWh，比 350 MWh 目标高 17 MWh；Harbor 则低于目标 12 MWh。"),
        ("h1", "3. 可再生能源与费率"),
        ("p", "园区光伏在本季度发电 184 MWh，占总实际用电 1,297 MWh 的 14.19%。财务估算统一使用 0.118 美元/千瓦时的购电费率；该费率用于比较，不包含需量电费和税费。"),
        ("pagebreak",),
        ("h1", "4. 改进措施与跨章节计算"),
        ("p", "Ridge 冷却塔变频改造计划在 2026 年 8 月 15 日前完成，预计每季度节省 22 MWh。即使完全实现该节省，也不能倒推 Q2 已达标，因为 Q2 是历史计量窗口，改造收益从 Q3 后半段开始记录。"),
        ("h1", "5. 干扰附注"),
        ("p", "Northwind 同名的 Harbor 仓储项目使用“Harbor-2”表计，不属于本报告。报告中的 Harbor 指园区主表 HBR-01；搜索到 Harbor-2 的 73 MWh 数据时不得加入本季度 1,297 MWh 合计。"),
    ])


MARKDOWN_DOCS = {
    "meridian-api-migration-v4.md": """# Meridian API Migration Guide v4.0

**Status:** Current  
**Effective:** 2026-07-01  
**v3 sunset:** 2026-11-30

## 1. Version authority

Version 4.0 is the current migration contract. Meridian v3 remains readable only during the transition window and is retired at 23:59 UTC on 2026-11-30. An archived v3 example later in this guide uses camelCase fields; it is intentionally retained as a migration contrast and must not be copied into new v4 clients.

## 2. Authentication

V4 uses OAuth 2.0 client credentials. Clients request a token from `/oauth2/token` with scope `meridian.write`; static `X-API-Key` authentication belongs to v3 and is rejected by v4 write endpoints. Tokens should be refreshed when fewer than 120 seconds remain.

### 2.1 Idempotency

Every create request must send `Idempotency-Key`. Meridian keeps the key and response binding for 24 hours. Reusing the same key with a different body returns HTTP 409; retrying the same body returns the original operation result.

## 3. Field mapping

| v3 field | v4 field | Rule |
|---|---|---|
| `customerId` | `customer_id` | required UUID |
| `completedAt` | `completed_at` | UTC RFC3339 timestamp |
| `orderItems` | `order_items` | array, max 200 |

The v4 completion timestamp is `completed_at` and must be UTC RFC3339, for example `2026-07-18T09:30:00Z`. A local time without an offset is invalid.

## 4. Rate limits and retry

The default tenant limit is 120 requests per minute with a burst of 30. Retry only HTTP 429 and 503, using delays of 200 ms, 400 ms, and 800 ms. Do not automatically retry HTTP 400 because it represents a caller correction, not transient capacity.

## 5. Connector dependency

Project Lumen release 3.4 requires Meridian v4 before rollout wave 2 begins. Wave 1 may observe v3 traffic, but wave 2 validation fails if the connector still sends `customerId` or `X-API-Key`.

## Appendix A - archived v3 contrast

V3 used `X-API-Key`, retained idempotency keys for 6 hours, and exposed `customerId`. These values are historical distractors, not valid v4 behavior.
""",
    "lumen-release-notes-3.4.md": """# Project Lumen Release Notes 3.4

Release 3.4 is the current production train for August 2026. It replaces 3.3 and requires database schema version 27 or later.

## Feature activation

Hybrid search is controlled by the exact feature flag `lumen.search.hybrid`. The similarly named `lumen.search.hybrid_preview` flag belonged to a closed beta and has no effect in 3.4.

## Rollout waves

| Wave | Date | Audience | Entry requirement |
|---|---|---|---|
| 1 | 2026-08-05 | internal tenants | schema 27 |
| 2 | 2026-08-12 | 10% external tenants | Meridian v4 connector |
| 3 | 2026-08-19 | 50% external tenants | error budget green |
| 4 | 2026-08-26 | all eligible tenants | change approval |

Wave 2 begins on 2026-08-12 and requires the Meridian v4 connector. This dependency is defined jointly with the Meridian migration guide, which retires v3 later in the year.

## Rollback threshold

Rollback the active wave when the five-minute error rate remains above 2.5% for 10 consecutive minutes. A single five-minute spike does not trigger rollback. The on-call lead records the start and end of the sustained window before disabling the flag.

## Cache behavior

Release 3.4 changes query-result cache TTL from 15 minutes to 5 minutes. The 15-minute value is the 3.3 default and is retained here only to explain the change. Authentication-token caching remains 2 minutes and is unrelated.

## Known non-blocking issue

Exported CSV filenames can omit the tenant display name when it contains emoji. The export data is complete, so this issue does not block rollout and does not change retrieval scoring.
""",
}


TXT_DOCS = {
    "orchard-inventory-procedure-2026.txt": """ORCHARD DISTRIBUTION - INVENTORY CONTROL PROCEDURE 2026
Status: CURRENT
Effective: 2026-02-01
Owner: Supply Chain Assurance

1. CYCLE COUNT SCHEDULE
Class A items are counted every Monday. Class B items are counted every Wednesday. Class C items are counted on the first business day of each calendar quarter. A holiday moves the count to the next business day; it does not cancel the count.

2. VARIANCE ACTIONS
CLASS      RECOUNT THRESHOLD      MANAGER REVIEW      FINANCE REVIEW
A          above 0.5%             above 1.0%          above 2.0%
B          above 1.0%             above 2.0%          above 4.0%
C          above 2.0%             above 4.0%          above 8.0%

For Class A inventory, a variance above 0.5% triggers an immediate recount, above 1.0% requires manager review, and above 2.0% additionally requires finance review. Thresholds are strictly “above”; exactly 0.5% does not trigger the recount rule.

3. COLD CHAIN
Refrigerated stock must remain between 2°C and 8°C. An excursion outside that range lasting more than 15 consecutive minutes requires quarantine under code QH-26 and a quality assessment. A brief sensor reading under 15 minutes is investigated but does not automatically require quarantine.

4. ROTATION TERMS
FEFO means first-expire, first-out and applies to dated food, medicine, and reagents. FIFO means first-in, first-out and applies only when expiration dates do not control disposition. The terms sound similar but FEFO takes precedence whenever an expiry date exists.

5. VERSION CONTROL
Code QH-26 replaces archived code QH-25. The 2025 procedure used a 3°C to 7°C target and different variance thresholds; those historical values must not be used for 2026 receipts.

6. DISTRACTOR NOTE
Parking permits, vehicle insurance carriers, and supplier contract duration are outside this inventory procedure. Their absence must not be filled with assumptions from unrelated warehouse manuals.
""",
    "orchard-inventory-archive-2025.txt": """ORCHARD DISTRIBUTION - INVENTORY CONTROL PROCEDURE 2025
Status: ARCHIVED - DO NOT USE FOR CURRENT RECEIPTS
Historical effective period: 2025-01-01 through 2026-01-31

1. HISTORICAL CYCLE COUNTS
Class A items were counted every Tuesday, Class B every second Thursday, and Class C twice per year. These schedules are preserved only for audit reconstruction.

2. HISTORICAL VARIANCE TABLE
CLASS      RECOUNT THRESHOLD      MANAGER REVIEW
A          above 0.8%             above 1.5%
B          above 1.5%             above 3.0%
C          above 3.0%             above 6.0%

3. HISTORICAL COLD CHAIN
The archived target range was 3°C to 7°C. Excursions longer than 20 minutes used quarantine code QH-25. On 2026-02-01 the current procedure widened the accepted range to 2°C through 8°C, shortened the duration threshold to 15 minutes, and replaced QH-25 with QH-26.

4. AUDIT WARNING
This archived file may rank highly for searches containing QH-25 or Tuesday. It is valid only for events inside its historical effective period and must never override the current 2026 procedure.

5. OUT-OF-SCOPE RECORDS
The archive does not name an insurance carrier, a parking level, or the duration of supplier audits. Those facts belong to other systems and are intentionally absent.
""",
}


def q(identifier, question, answer, sources, evidence, tags, conversation=None, turn=1):
    return {
        "id": identifier,
        "question": question,
        "reference_answer": answer,
        "expected_sources": sources,
        "expected_evidence": evidence,
        "tags": tags,
        "conversation_group": conversation or identifier,
        "turn": turn,
        "unanswerable": "unanswerable" in tags,
    }


def questions():
    ao = "aurora-transit-operations-v2.1.pdf"
    ab = "aurora-maintenance-bulletin-2026-rB.pdf"
    ne = "northwind-energy-q2-2026.pdf"
    hc = "helios-travel-policy-v3.2.docx"
    ha = "helios-travel-policy-v2.8-archived.docx"
    at = "atlas-support-playbook-2026.4.docx"
    me = "meridian-api-migration-v4.md"
    lu = "lumen-release-notes-3.4.md"
    oc = "orchard-inventory-procedure-2026.txt"
    oa = "orchard-inventory-archive-2025.txt"
    return [
        q("RW-001", "Aurora 运营手册当前版本何时生效，替代哪个版本？", "版本 2.1 于 2026-04-01 生效，替代版本 1.9。", [ao], ["版本 2.1 自 2026 年 4 月 1 日生效，替代版本 1.9。"], ["direct_fact", "pdf"]),
        q("RW-002", "P1 服务关键事件的确认和接驳时限分别是多少？", "10 分钟内确认，25 分钟内安排替代接驳车。", [ao], ["P1 服务关键事件要求控制中心在 10 分钟内确认，并在 25 分钟内安排替代接驳车。"], ["cross_section", "pdf"]),
        q("RW-003", "机场快线最低应保留多少比例的备用车辆？", "15%。", [ao], ["机场快线的高峰计划间隔为 15 分钟，非高峰为 20 分钟，最低备用车辆比例为 15%。"], ["table_numeric", "pdf"]),
        q("RW-004", "安全停车和服务暂停是不是同一个等级？", "不是。安全停车是基于安全风险的强制停止并对应 P0；服务暂停是计划性调度动作，本身不等于 P0。", [ao], ["“Safety Stop 安全停车”是基于安全风险的强制停止，对应 P0；“Service Pause 服务暂停”是为恢复班次或调度间隔而进行的短时计划动作，本身不等于 P0。"], ["semantic_paraphrase", "similar_term_distractor", "pdf"]),
        q("RW-005", "损坏锂电池失物应该保管 90 天吗？", "不应该，应在 24 小时内转交安全团队。", [ao], ["含裸露或损坏锂电池的物品属于危险品，必须在 24 小时内转交安全团队，不能按普通失物保管 90 天。"], ["similar_content_distractor", "pdf"]),
        q("RW-006", "什么组合条件会触发高架区段暂停，暂停后还要做什么？", "暴雨红色预警且连续 10 分钟平均风速超过 24 米/秒；暂停后 15 分钟内发布乘客公告并协调地面接驳。", [ao], ["当气象台发布暴雨红色预警，且高架区段连续 10 分钟测得平均风速超过 24 米/秒时，值班主任应暂停受影响高架区段", "暂停后还必须在 15 分钟内发布乘客公告，并协调地面接驳"], ["cross_section", "pdf"]),
        q("RW-007", "Revision B 要求的 NX-7 执行器安装扭矩是多少？", "42 N·m。", [ab], ["它把 NX-7 车门执行器安装扭矩从 Revision A 的 38 N·m 更正为 42 N·m"], ["version_conflict", "pdf"]),
        q("RW-008", "NX-7 车门执行器按什么周期检查？", "每 18,000 公里或每 120 天，以先到者为准。", [ab], ["NX-7 车门执行器应每 18,000 公里或每 120 天检查一次，以先到者为准。"], ["semantic_paraphrase", "pdf"]),
        q("RW-009", "AX-440 和 AX-404 分别是什么零件？", "AX-440 是导向滚轮，AX-404 是位置传感器。", [ab], ["导向滚轮零件号是 AX-440，位置传感器零件号是 AX-404；两个编号数字相似，但不得互换。"], ["similar_term_distractor", "pdf"]),
        q("RW-010", "Central Depot 的停电窗口和车辆范围是什么？", "2026-06-08 01:10-03:40，覆盖 NX-7 021-036。", [ab], ["Central Depot 的维护停电窗口为 2026 年 6 月 8 日 01:10 至 03:40，覆盖 NX-7 021 至 036。"], ["table_numeric", "pdf"]),
        q("RW-011", "同一车厢两个相邻车门不可用时按什么等级处理，确认和接驳时限是什么？", "升级为 P1；10 分钟内确认，25 分钟内安排接驳。", [ab, ao], ["同一车厢两个相邻车门不可用时，应升级为 P1", "P1 服务关键事件要求控制中心在 10 分钟内确认，并在 25 分钟内安排替代接驳车。"], ["cross_document", "pdf"]),
        q("RW-012", "工单关闭前必须提交哪三类证据？", "扭矩扳手读数照片、执行器序列号、双人复核签名。", [ab], ["关闭工单前必须上传扭矩扳手读数照片、执行器序列号和双人复核签名。"], ["direct_fact", "pdf"]),
        q("RW-013", "Northwind Q2 三个站点实际用电合计多少？", "1,297 MWh。", [ne], ["三个站点实际用电合计 1,297 MWh"], ["table_numeric", "pdf"]),
        q("RW-014", "Ridge 实际用电比目标高多少？", "高 17 MWh。", [ne], ["Ridge 实际用电 367 MWh，比 350 MWh 目标高 17 MWh"], ["table_numeric", "pdf"]),
        q("RW-015", "园区光伏发电量及其占总用电比例是多少？", "184 MWh，占 14.19%。", [ne], ["园区光伏在本季度发电 184 MWh，占总实际用电 1,297 MWh 的 14.19%。"], ["cross_section", "table_numeric", "pdf"]),
        q("RW-016", "Harbor 最终实际用电是 421 MWh 吗？", "不是，421 MWh 是 provisional 值，Final 值为 418 MWh。", [ne], ["provisional 工作表曾把 Harbor 站点记录为 421 MWh，但最终对账确认其实际用电为 418 MWh"], ["version_conflict", "pdf"]),
        q("RW-017", "报告比较购电成本时使用什么统一费率？", "0.118 美元/千瓦时。", [ne], ["财务估算统一使用 0.118 美元/千瓦时的购电费率"], ["direct_fact", "pdf"]),
        q("RW-018", "Q2 合计超目标多少 MWh 和百分比？", "17 MWh，1.33%。", [ne], ["目标合计 1,280 MWh，超出 17 MWh，即 1.33%。"], ["cross_section", "table_numeric", "pdf"]),
        q("RW-019", "Helios 当前差旅政策何时生效，并替代哪个版本？", "v3.2 于 2026-06-01 生效，替代 v2.8。", [hc], ["版本 3.2 自 2026 年 6 月 1 日生效，并明确替代版本 2.8"], ["direct_fact", "version_conflict", "docx"]),
        q("RW-020", "按当前政策，单程飞行 9 小时可以订什么舱等？", "可预订高级经济舱。", [hc], ["6 小时至 10 小时可预订高级经济舱"], ["direct_fact", "contextual_followup", "docx"], "ctx-travel", 1),
        q("RW-021", "那如果同时参加上海指定会议，酒店例外上限是多少？", "每晚 1,080 元。", [hc], ["上海的会议例外上限为每晚 1,080 元"], ["table_numeric", "contextual_followup", "docx"], "ctx-travel", 2),
        q("RW-022", "境内餐费和必要杂费每天分别封顶多少？", "餐费 300 元，必要杂费 80 元。", [hc], ["境内差旅餐费按每日 300 元封顶，必要杂费按每日 80 元封顶。"], ["direct_fact", "docx"]),
        q("RW-023", "机票应提前多久预订？时间例外是否自动允许商务舱？", "至少提前 14 个自然日；时间窗口例外不自动构成商务舱授权。", [hc], ["所有机票原则上应在出发前至少 14 个自然日完成预订", "时间窗口例外与舱等例外是两个独立审批事项"], ["cross_section", "docx"]),
        q("RW-024", "当前与归档政策对 9 小时航程的商务舱规则有何冲突，当前应采用哪条？", "归档 v2.8 对超过 8 小时可由部门负责人批准商务舱；当前 v3.2 需超过 10 小时且业务副总裁书面批准，所以 9 小时当前只能高级经济舱。", [hc, ha], ["超过 10 小时仅在业务副总裁书面批准后可预订商务舱", "超过 8 小时在部门负责人批准后可使用商务舱"], ["cross_document", "version_conflict", "docx"]),
        q("RW-025", "归档 v2.8 的上海住宿上限是多少？", "每晚 760 元。", [ha], ["归档版本中，上海住宿上限为每晚 760 元"], ["direct_fact", "docx"]),
        q("RW-026", "归档政策下 9 小时航程由谁批准商务舱？", "部门负责人。", [ha], ["超过 8 小时在部门负责人批准后可使用商务舱"], ["semantic_paraphrase", "docx"]),
        q("RW-027", "为什么 v2.8 不能用于 2026-06-01 之后的新预订？", "因为它已被 v3.2 替代，只能用于历史报销审计。", [ha, hc], ["本文件自 2026 年 6 月 1 日起不再有效", "归档政策只能用于历史报销审计，不能用于新的预订或审批"], ["cross_document", "version_conflict", "docx"]),
        q("RW-028", "v2.8 的历史提前预订窗口是多少？", "至少 7 个自然日。", [ha], ["机票应至少提前 7 个自然日预订。"], ["direct_fact", "docx"]),
        q("RW-029", "归档版本境内餐费每天是多少？", "260 元。", [ha], ["境内餐费为每日 260 元"], ["table_numeric", "docx"]),
        q("RW-030", "上海普通住宿上限从归档版到当前版增加了多少？", "增加 140 元，从 760 元到 900 元。", [ha, hc], ["上海普通住宿上限后来增加了 140 元", "上海普通住宿上限为每晚 900 元"], ["cross_document", "table_numeric", "docx"]),
        q("RW-031", "S1 的首次响应和更新频率是什么？", "15 分钟内首次响应，之后每 30 分钟更新。", [at], ["S1 事件的首次响应目标是 15 分钟，随后每 30 分钟更新一次"], ["table_numeric", "contextual_followup", "docx"], "ctx-atlas", 1),
        q("RW-032", "P1 是否必然等于 S1？", "不是。P1 是内部优先级，S1 是客户影响严重度；S2、P1 的组合也可能存在。", [at], ["一个影响单个战略客户但有替代方案的问题可以是 S2、P1；因此客服不得把 P1 自动说成 S1"], ["similar_term_distractor", "docx"]),
        q("RW-033", "那么确认这种事件后，高管最迟何时通知？", "确认 S1 后 20 分钟内。", [at], ["高管通知必须在确认 S1 后 20 分钟内完成。"], ["contextual_followup", "docx"], "ctx-atlas", 2),
        q("RW-034", "工程修复完成后，哪三项动作完成才能宣布恢复？", "客户沟通负责人确认外部状态页更新、记录员链接时间线、事件指挥官宣布恢复。", [at], ["客户沟通负责人必须确认外部状态页已更新，记录员必须链接时间线，最后由事件指挥官宣布恢复。"], ["cross_section", "docx"]),
        q("RW-035", "S2 的首次响应和更新间隔是多少？", "1 小时首次响应，每 2 小时更新。", [at], ["S2 的首次响应目标是 1 小时，更新间隔为 2 小时。"], ["table_numeric", "docx"]),
        q("RW-036", "Atlas 手册规定的客户退款比例上限是多少？", "文档没有规定退款金额、信用额度或赔偿上限，无法回答具体比例。", [], [], ["unanswerable", "docx"]),
        q("RW-037", "Meridian v3 在什么时候停止服务？", "2026-11-30 23:59 UTC。", [me], ["retired at 23:59 UTC on 2026-11-30"], ["direct_fact", "contextual_followup", "markdown"], "ctx-meridian", 1),
        q("RW-038", "迁移到新版本后，它使用什么鉴权，旧方式还能写入吗？", "v4 使用 OAuth 2.0 client credentials；v3 的静态 X-API-Key 会被 v4 写端点拒绝。", [me], ["V4 uses OAuth 2.0 client credentials", "static `X-API-Key` authentication belongs to v3 and is rejected by v4 write endpoints"], ["contextual_followup", "version_conflict", "markdown"], "ctx-meridian", 2),
        q("RW-039", "v4 幂等键保留多久？相同键换请求体会怎样？", "保留 24 小时；相同键配不同请求体返回 HTTP 409。", [me], ["Meridian keeps the key and response binding for 24 hours. Reusing the same key with a different body returns HTTP 409"], ["cross_section", "markdown"]),
        q("RW-040", "哪些状态允许自动重试，退避序列是什么？", "只重试 HTTP 429 和 503，延迟依次为 200、400、800 毫秒。", [me], ["Retry only HTTP 429 and 503, using delays of 200 ms, 400 ms, and 800 ms."], ["direct_fact", "markdown"]),
        q("RW-041", "v3 的 customerId 在 v4 改成什么，格式要求是什么？", "改为 `customer_id`，要求 UUID。", [me], ["`customerId` | `customer_id` | required UUID"], ["table_numeric", "semantic_paraphrase", "markdown"]),
        q("RW-042", "completed_at 必须采用哪种时间格式？", "UTC RFC3339，例如 2026-07-18T09:30:00Z。", [me], ["must be UTC RFC3339, for example `2026-07-18T09:30:00Z`"], ["direct_fact", "markdown"]),
        q("RW-043", "Lumen 3.4 启用混合搜索的准确 feature flag 是什么？", "`lumen.search.hybrid`。", [lu], ["Hybrid search is controlled by the exact feature flag `lumen.search.hybrid`."], ["similar_term_distractor", "markdown"]),
        q("RW-044", "Lumen 3.4 要求的最低数据库 schema 版本是多少？", "27。", [lu], ["requires database schema version 27 or later"], ["direct_fact", "markdown"]),
        q("RW-045", "什么错误率条件会触发当前波次回滚？", "五分钟错误率持续高于 2.5%，并连续保持 10 分钟。", [lu], ["when the five-minute error rate remains above 2.5% for 10 consecutive minutes"], ["cross_section", "markdown"]),
        q("RW-046", "3.4 的查询结果缓存 TTL 是多少，15 分钟代表什么？", "3.4 为 5 分钟；15 分钟是 3.3 的旧默认值。", [lu], ["Release 3.4 changes query-result cache TTL from 15 minutes to 5 minutes. The 15-minute value is the 3.3 default"], ["version_conflict", "similar_content_distractor", "markdown"]),
        q("RW-047", "Lumen 第二波何时开始，依赖 Meridian 哪个版本及鉴权方式？", "第二波 2026-08-12 开始，要求 Meridian v4，并使用 OAuth 2.0 client credentials。", [lu, me], ["Wave 2 begins on 2026-08-12 and requires the Meridian v4 connector", "V4 uses OAuth 2.0 client credentials"], ["cross_document", "table_numeric", "markdown"]),
        q("RW-048", "Lumen 第三波面向多少外部租户，进入条件是什么？", "50% 外部租户，要求 error budget green。", [lu], ["3 | 2026-08-19 | 50% external tenants | error budget green"], ["table_numeric", "markdown"]),
        q("RW-049", "Orchard 当前程序中 A 类和 B 类分别哪天盘点？", "A 类每周一，B 类每周三。", [oc], ["Class A items are counted every Monday. Class B items are counted every Wednesday."], ["direct_fact", "contextual_followup", "txt"], "ctx-orchard", 1),
        q("RW-050", "那 A 类差异达到什么条件要立即复盘、经理复核和财务复核？", "高于 0.5% 立即复盘，高于 1.0% 经理复核，高于 2.0% 财务复核。", [oc], ["above 0.5% triggers an immediate recount, above 1.0% requires manager review, and above 2.0% additionally requires finance review"], ["table_numeric", "contextual_followup", "txt"], "ctx-orchard", 2),
        q("RW-051", "当前冷链温度范围和强制隔离的持续时长条件是什么？", "2°C 到 8°C；超出范围持续超过 15 分钟时强制隔离。", [oc], ["must remain between 2°C and 8°C. An excursion outside that range lasting more than 15 consecutive minutes requires quarantine"], ["cross_section", "txt"]),
        q("RW-052", "有有效期的商品应使用 FIFO 还是 FEFO，为什么？", "使用 FEFO，因为有有效期时先到期先出优先于先进先出。", [oc], ["FEFO means first-expire, first-out", "FEFO takes precedence whenever an expiry date exists"], ["semantic_paraphrase", "similar_term_distractor", "txt"]),
        q("RW-053", "2026 当前隔离代码是什么，它替代了什么？", "QH-26，替代归档代码 QH-25。", [oc, oa], ["Code QH-26 replaces archived code QH-25."], ["cross_document", "version_conflict", "txt"]),
        q("RW-054", "当前程序要求供应商审计持续多少天？", "文档没有提供供应商审计持续时间，无法回答。", [], [], ["unanswerable", "txt"]),
        q("RW-055", "2025 归档程序的冷链目标范围是多少？", "3°C 到 7°C。", [oa], ["The archived target range was 3°C to 7°C."], ["direct_fact", "txt"]),
        q("RW-056", "2025 归档隔离代码和时长门槛是什么？", "QH-25，超出范围超过 20 分钟。", [oa], ["Excursions longer than 20 minutes used quarantine code QH-25."], ["direct_fact", "txt"]),
        q("RW-057", "为什么不能把每周二作为 2026 年 A 类盘点日？", "每周二是 2025 归档规则；2026 当前规则为每周一。", [oa, oc], ["Class A items were counted every Tuesday", "Class A items are counted every Monday."], ["cross_document", "version_conflict", "txt"]),
        q("RW-058", "从 2025 到 2026，冷链范围、时长门槛和隔离代码如何变化？", "从 3-7°C、超过 20 分钟、QH-25，改为 2-8°C、超过 15 分钟、QH-26。", [oa, oc], ["widened the accepted range to 2°C through 8°C, shortened the duration threshold to 15 minutes, and replaced QH-25 with QH-26"], ["cross_document", "version_conflict", "txt"]),
        q("RW-059", "Orchard 使用哪家车辆保险公司？", "这些文档没有提供保险公司信息，无法回答。", [], [], ["unanswerable", "txt"]),
        q("RW-060", "Northwind Harbor 园区员工停车在第几层？", "文档没有停车楼层信息，无法回答。", [], [], ["unanswerable", "pdf"]),
    ]


def main():
    CORPUS.mkdir(parents=True, exist_ok=True)
    build_helios_current(CORPUS / "helios-travel-policy-v3.2.docx")
    build_helios_archived(CORPUS / "helios-travel-policy-v2.8-archived.docx")
    build_atlas(CORPUS / "atlas-support-playbook-2026.4.docx")
    build_pdfs()
    for filename, content in MARKDOWN_DOCS.items():
        (CORPUS / filename).write_text(content.strip() + "\n", encoding="utf-8")
    for filename, content in TXT_DOCS.items():
        (CORPUS / filename).write_text(content.strip() + "\n", encoding="utf-8")
    items = questions()
    if len(items) != 60:
        raise RuntimeError(f"expected 60 questions, got {len(items)}")
    with DATASET.open("w", encoding="utf-8", newline="\n") as stream:
        for item in items:
            stream.write(json.dumps(item, ensure_ascii=False) + "\n")
    files = sorted(CORPUS.iterdir())
    manifest = {
        "name": "real-world-v1",
        "purpose": "真实格式、真实 Provider、公开 API 端到端评测语料",
        "document_count": len(files),
        "question_count": len(items),
        "documents": [{"filename": item.name, "format": item.suffix.lower().lstrip("."), "size_bytes": item.stat().st_size} for item in files],
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"documents": len(files), "questions": len(items), "corpus": str(CORPUS)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
